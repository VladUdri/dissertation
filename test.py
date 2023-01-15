from time import sleep
import win32com.client as comctl
import mouse
import subprocess
import docx
from docx.enum.text import WD_COLOR_INDEX
import os
import numpy as nm

import pytesseract

# importing OpenCV
import cv2

from PIL import ImageGrab
from docx_commands import DocxCommands

# doc = docx.Document('gfg.docx')
# doc.add_heading('GeeksForGeeks', 0)
# para = doc.add_paragraph('''GeeksforGeeks is a Computer Science portal for geeks.''')
# para.add_run(''' It contains well written, well thought and well-explained '''
#              ).font.highlight_color = WD_COLOR_INDEX.YELLOW
# para.add_run('''computer science and programming articles, quizzes etc.''')
#
# doc.save('gfg.docx')


# wsh = comctl.Dispatch("WScript.Shell")
# Google Chrome window title
# try:

# mouse.move("600", "200")
# mouse.double_click()
# wsh.AppActivate("Chrome")
# sleep(0.5)
# wsh.SendKeys("^{Esc}")
# sleep(0.5)
# wsh.SendKeys("word")
# sleep(0.5)
# wsh.SendKeys("{ENTER}")
# sleep(5)
# wsh.SendKeys("{ENTER}")
# sleep(4)
#
# wsh.SendKeys("This is only a test.")
#
# finally:
#     print("done")

# if __name__ == "__main__":
#     # doc_name = input('Doc name: ')
#     # command = input('Command: ')
#     # doc_name = 'test'
#     # command = 'paragraph'
#     # # print(input1)
#     # doc = DocxCommands(name_of_doc=doc_name, command=command)
#     # doc.runCommand()
#
#     """
#     Steps:
#     - User the command to open document/create new document. For both of them he has to specify the name
#     - User is asked to provide the formatting of the document
#         The user can say color, font, alignment, font size etc. (TODO add mode options later)
#     - User can add paragraph, title, headline etc.
#         Commands like: "Add title", "Add new paragraph", "Continue"
#     - After adding some text, the user will be able to specify if he wants a new format for that part.
#
#
#     - What commands should I add, right now?
#     - Work on automations and user preferences (I like to use Calibri and my font size to be 11, with my title being 14.)
#
#
#     """

def imToString():

	# Path of tesseract executable
	pytesseract.pytesseract.tesseract_cmd ='D:\\pythonProject1\\venv\\Lib\\site-packages\\pytesseract'
	while(True):

		# ImageGrab-To capture the screen image in a loop.
		# Bbox used to capture a specific area.
		cap = ImageGrab.grab(bbox =(700, 300, 1400, 900))

		# Converted the image to monochrome for it to be easily
		# read by the OCR and obtained the output String.
		tesstr = pytesseract.image_to_string(
				cv2.cvtColor(nm.array(cap), cv2.COLOR_BGR2GRAY),
				lang ='eng')
		print(tesstr)

# Calling the function
imToString()
