import docx
from docx.enum.text import WD_COLOR_INDEX
import os
import pathlib
from time import sleep
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

blue = RGBColor(0x00, 0x00, 0xFF)
red = RGBColor(0xFF, 0x00, 0x00)
green = RGBColor(0x00, 0xFF, 0x00)
black = RGBColor(0x00, 0x00, 0x00)

arial = 'Arial'
times_new_roman = 'Times New Roman'
calibri = 'Calibri'


class DocxCommands:
    def __init__(self, name_of_doc, command):
        self.name_of_doc = name_of_doc
        self.command = command
        self.name_of_doc = self.addFileType(self.name_of_doc)

    def addFileType(self, name_of_doc):
        if name_of_doc.find('.docx') == -1:
            name_of_doc = name_of_doc + '.docx'
        return name_of_doc

    def getDoc(self):
        doc = docx.Document(self.name_of_doc)
        print(doc)

    def find(self, name, path):
        for root, dirs, files in os.walk(path):
            if name in files:
                return os.path.join(root, name)

    def initDoc(self):
        path = pathlib.Path().resolve()
        if self.find(self.name_of_doc, path) is None:
            newDocx = docx.Document()
            newDocx.save(self.name_of_doc)
        else:
            newDocx = docx.Document(self.name_of_doc)

        return newDocx

    def addHeadline(self, doc):
        headline_text = input('Headline Text: ')
        doc.add_heading(headline_text, 0)
        doc.save(self.name_of_doc)

    def addParagraph(self, doc):
        para_font = doc.add_paragraph()
        paragraph = para_font.add_run(
            "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.")
        options = input('Format options: ')
        format_options = options.split(' ')
        for format_option in format_options:
            print(format_option)
            self.formatParagraph(paragraph=paragraph, para_font=para_font, format_option=format_option)
        doc.save(self.name_of_doc)
        print('Done!')

    def formatParagraph(self, paragraph, para_font, format_option):
        value = self.getValue(format_option)
        if format_option == 'size':
            self.formatSize(paragraph, value)
        elif format_option == 'font':
            self.formatFont(paragraph, value)
        elif format_option == 'color':
            self.formatColor(paragraph, value)
        elif format_option == 'style':
            self.formatFontStyle(paragraph, value)
        elif format_option == 'align':
            self.formatAlignment(para_font, value)

    def formatSize(self, paragraph, size):
        paragraph.font.size = Pt(int(size))

    def formatColor(self, paragraph, color):
        if color == 'red':
            paragraph.font.color.rgb = red
        elif color == 'blue':
            paragraph.font.color.rgb = blue
        elif color == 'green':
            paragraph.font.color.rgb = green
        elif color == 'black':
            paragraph.font.color.rgb = black

    def formatAlignment(self, paragraph, align):
        if align == 'center':
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align == 'left':
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align == 'right':
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align == 'justify':
            paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def formatFont(self, paragraph, font):
        if font == 'arial':
            paragraph.font.name = arial
        elif font == 'calibri':
            paragraph.font.name = calibri
        elif font == 'times' or font == 'times new' or font == 'times new roman':
            paragraph.font.name = times_new_roman

    def formatFontStyle(self, paragraph, font_style):
        if font_style == 'bold':
            paragraph.font.bold = True
        elif font_style == 'italic':
            paragraph.font.italic = True
        elif font_style == 'underline':
            paragraph.font.underline = True

    def getValue(self, type):
        value = input('Please enter ' + type + ': ')
        return value

    def runCommand(self):
        doc = self.initDoc()
        if self.command == 'headline':
            self.addHeadline(doc)
        if self.command == 'paragraph':
            self.addParagraph(doc)

    # sleep(1)
    # os.system('D:\\pythonProject1\\test.docx')
